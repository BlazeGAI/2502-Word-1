import streamlit as st
import docx
import re
import pandas as pd
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

def check_word_document(doc):
    checklist_data = {
        "Grading Criteria": [
            "1. Is the font Times New Roman, 12pt?",
            "2. Is line spacing set to double?",
            "3. Are margins set to 1 inch on all sides?",
            "4. Is the title on the title page centered and bold?",
            "5. Is the title on the second page not bold?",
            "6. Are the paragraphs left-aligned?",
            "7. Are there at least 3 paragraphs?",
            "8. Is there a References page?",
            "9. Are in-text citations properly formatted?",
            "10. Is there a title page?",
            "11. Is the title page center aligned?",
            "12. Are there page numbers in the top right?"
        ],
        "Completed": []
    }

    # Ensure all checks append a result
    def safe_append(value):
        checklist_data["Completed"].append("Yes" if value else "No")

    # Retrieve document default font settings
    default_font_name = doc.styles['Normal'].font.name if 'Normal' in doc.styles else "Times New Roman"
    default_font_size = doc.styles['Normal'].font.size if 'Normal' in doc.styles else Pt(12)
    assumed_font_size = default_font_size or Pt(12)

    def is_correct_font(paragraph):
        for run in paragraph.runs:
            run_font = run.font.name or default_font_name
            run_size = run.font.size or assumed_font_size
            
            if run_font != 'Times New Roman' or run_size != Pt(12):
                return False
        return True
    
    safe_append(all(is_correct_font(p) for p in doc.paragraphs if p.text.strip()) if doc.paragraphs else False)
    safe_append(all(
        p.paragraph_format.line_spacing in [None, 2.0] for p in doc.paragraphs if p.text.strip()
    ) if doc.paragraphs else False)
    safe_append(all(
        section.left_margin.inches == 1 and
        section.right_margin.inches == 1 and
        section.top_margin.inches == 1 and
        section.bottom_margin.inches == 1
        for section in doc.sections
    ) if doc.sections else False)

    first_non_empty_paragraph = next((p for p in doc.paragraphs if p.text.strip()), None)
    safe_append(
        first_non_empty_paragraph and 
        first_non_empty_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and
        any(run.bold for run in first_non_empty_paragraph.runs)
    )

    second_page_title = doc.paragraphs[1] if len(doc.paragraphs) > 1 else None
    safe_append(second_page_title and not any(run.bold for run in second_page_title.runs))

    body_paragraphs = []
    found_references = False
    header_done = False
    
    for p in doc.paragraphs:
        text = p.text.strip() if p.text else ""
        if not text:
            continue
        if not header_done and len(text) > 100:
            header_done = True
        if text.lower() == 'references':
            found_references = True
            continue
        if header_done and not found_references and len(text) > 100:
            body_paragraphs.append(p)

    body_text_paragraphs = [p for p in doc.paragraphs if len(p.text.strip()) > 50] if doc.paragraphs else []
    safe_append(all(p.alignment in [WD_ALIGN_PARAGRAPH.LEFT, None] for p in body_text_paragraphs) if body_text_paragraphs else False)
    safe_append(len(body_paragraphs) >= 3)
    safe_append(any(p.text.strip().lower() == 'references' for p in doc.paragraphs) if doc.paragraphs else False)
    safe_append(any(
        found_references and '(' in p.text and ')' in p.text for p in doc.paragraphs
    ) if doc.paragraphs else False)

    citation_pattern = re.compile(r'\. \([A-Za-z]+, \d{4}\)')
    safe_append(any(
        citation_pattern.search(p.text if p.text else "") for p in body_paragraphs
    ) if body_paragraphs else False)

    title_page_paragraphs = [p for p in doc.paragraphs[:5] if p.text.strip()]
    safe_append(len(title_page_paragraphs) > 0 and all(len(p.text) < 100 for p in title_page_paragraphs))
    safe_append(
        title_page_paragraphs and all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in title_page_paragraphs)
    )

    # Improved page number detection: Ensure every section's header has a right-aligned numeral
    has_page_numbers = all(
        section.header and any(
            para.alignment == WD_ALIGN_PARAGRAPH.RIGHT and para.text.strip().isdigit()
            for para in section.header.paragraphs if para.text.strip()
        ) for section in doc.sections
    ) if doc.sections else False
    safe_append(has_page_numbers)

    # Ensure both lists have the same length
    while len(checklist_data["Completed"]) < len(checklist_data["Grading Criteria"]):
        checklist_data["Completed"].append("No")
    while len(checklist_data["Completed"]) > len(checklist_data["Grading Criteria"]):
        checklist_data["Grading Criteria"].append(f"{len(checklist_data["Grading Criteria"])+1}. Unknown Criteria")

    return checklist_data

def display_results(checklist_data):
    # Calculate scores
    total_yes = checklist_data["Completed"].count("Yes")
    total_items = len(checklist_data["Completed"])
    percentage_complete = (total_yes / total_items) * 100
    points = (total_yes / total_items) * 20

    # Display scores
    col1, col2 = st.columns(2)
    
    with col1:
        if percentage_complete == 100:
            st.success(f"Completion Score: {percentage_complete:.1f}%")
        elif percentage_complete >= 80:
            st.warning(f"Completion Score: {percentage_complete:.1f}%")
        else:
            st.error(f"Completion Score: {percentage_complete:.1f}%")

    with col2:
        if points == 20:
            st.success(f"Points: {points:.1f}/20")
        elif points >= 16:
            st.warning(f"Points: {points:.1f}/20")
        else:
            st.error(f"Points: {points:.1f}/20")

    # Display checklist
    st.subheader("Detailed Checklist")
    checklist_df = pd.DataFrame(checklist_data)
    st.table(checklist_df)

st.title("Word Document Grading Checklist")

uploaded_file = st.file_uploader("Upload a .docx file", type=["docx"])

if uploaded_file:
    doc = docx.Document(uploaded_file)
    results = check_word_document(doc)
    display_results(results)
